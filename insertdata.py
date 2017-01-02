from docx import Document
import os
from pymongo import MongoClient


def get_data():
    '''
    Function to crawl through .docx files in the /static directory
    and insert data into a MongoDB database.
    '''

    data_for_insertion = []
    for root, dirs, files in os.walk("./static"):
        for file in files:
            if file.endswith(".docx"):
                f = open(os.path.join(root, file), 'rb')
                document = Document(f)
                f.close()
                keys = []
                values = []
                i = 0
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                i += 1
                                if i % 2 != 0:
                                    keys.append(paragraph.text)
                                else:
                                    values.append(paragraph.text)
                data_dictionary = dict(zip(keys, values))
                data_for_insertion.append(data_dictionary)

    client = MongoClient()
    db = client.pat_dose_ests_db
    pat_dose_ests = db.pat_dose_ests
    pat_dose_ests.insert_many(data_for_insertion)

    print("There are %s documents in the collection" % (pat_dose_ests.count()))

if __name__ == "__main__":
    get_data()
